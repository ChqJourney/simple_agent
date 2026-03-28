from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterable

try:
    from docx import Document as DocxDocument  # type: ignore
    from docx.oxml.table import CT_Tbl  # type: ignore
    from docx.oxml.text.paragraph import CT_P  # type: ignore
    from docx.table import Table as DocxTable  # type: ignore
    from docx.text.paragraph import Paragraph as DocxParagraph  # type: ignore
except Exception:  # pragma: no cover - exercised via runtime validation
    DocxDocument = None
    CT_Tbl = None
    CT_P = None
    DocxTable = None
    DocxParagraph = None


MAX_WORD_FILE_BYTES = 10 * 1024 * 1024
MAX_TITLE_CHARS = 80
HEADING_STYLE_RE = re.compile(r"Heading\s+(?P<level>\d+)", re.IGNORECASE)
NUMBERED_HEADING_RE = re.compile(r"^(?P<number>\d+(?:\.\d+){0,5})[\s.)-]+(?P<title>.+?)\s*$")
TRF_HEADER_KEYWORDS = {
    "clause",
    "requirement",
    "verdict",
    "result",
    "remark",
    "remarks",
    "comment",
    "comments",
    "evidence",
    "finding",
    "observation",
    "decision",
    "test",
    "inspection",
}


def _require_python_docx() -> None:
    if DocxDocument is None:
        raise RuntimeError("python-docx is not installed. Install `python-docx` to use Word document tools.")


def _normalize_text(text: Any) -> str:
    return " ".join(str(text or "").split()).strip()


def _make_anchor(text: str) -> str:
    anchor = re.sub(r"[^\w\s-]", "", text.lower()).strip()
    anchor = re.sub(r"[\s_]+", "-", anchor)
    return anchor


def _truncate_title(text: str) -> str:
    normalized = _normalize_text(text)
    if len(normalized) <= MAX_TITLE_CHARS:
        return normalized
    return normalized[:MAX_TITLE_CHARS].rstrip() + "..."


def _paragraph_level(text: str, style_name: str) -> int | None:
    style_match = HEADING_STYLE_RE.search(style_name)
    if style_match:
        return max(1, int(style_match.group("level")))

    numbered_match = NUMBERED_HEADING_RE.match(text)
    if numbered_match:
        return numbered_match.group("number").count(".") + 1

    return None


def _iter_block_items(document: Any) -> Iterable[tuple[str, Any]]:
    body = getattr(getattr(document, "element", None), "body", None)
    if body is None or CT_P is None or CT_Tbl is None or DocxParagraph is None or DocxTable is None:
        for paragraph in getattr(document, "paragraphs", []):
            yield "paragraph", paragraph
        for table in getattr(document, "tables", []):
            yield "table", table
        return

    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", DocxParagraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield "table", DocxTable(child, document)


def _format_row_text(cells: list[dict[str, Any]], table_type: str = "grid") -> str:
    values = [str(cell.get("text") or "").strip() for cell in cells if str(cell.get("text") or "").strip()]
    if not values:
        return ""
    if table_type == "key_value" and len(values) == 2 and len(values[0]) <= 80:
        return f"{values[0]}: {values[1]}"
    return " | ".join(values)


def _looks_like_caption(text: str) -> bool:
    lowered = _normalize_text(text).lower()
    return lowered.startswith("table ") or lowered.startswith("表 ")


class WordReader:
    def __init__(self, docx_path: str | Path):
        _require_python_docx()
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"Word file not found: {self.docx_path}")
        file_size = self.docx_path.stat().st_size
        if file_size > MAX_WORD_FILE_BYTES:
            raise ValueError(f"File too large: {file_size} bytes (max: {MAX_WORD_FILE_BYTES} bytes)")

        self.doc = DocxDocument(self.docx_path)
        self.blocks: list[dict[str, Any]] = []
        self.paragraphs: list[dict[str, Any]] = []
        self.tables: list[dict[str, Any]] = []
        self._collect_blocks()

    def _collect_blocks(self) -> None:
        paragraph_index = 0
        active_heading: dict[str, Any] | None = None
        last_non_empty_paragraph: str | None = None

        for block_type, block in _iter_block_items(self.doc):
            if block_type == "paragraph":
                paragraph_index += 1
                item = self._serialize_paragraph(block, paragraph_index)
                self.paragraphs.append(item)
                self.blocks.append({"block_type": "paragraph", **item})
                if item["text"]:
                    last_non_empty_paragraph = item["text"]
                if item["text"] and item["level"] is not None:
                    active_heading = {
                        "title": item["text"],
                        "level": int(item["level"]),
                        "paragraph_index": item["paragraph_index"],
                    }
                continue

            table_index = len(self.tables) + 1
            table_item = self._serialize_table(
                block,
                table_index=table_index,
                preceding_heading=active_heading,
                preceding_paragraph=last_non_empty_paragraph,
                preceding_paragraph_index=paragraph_index or None,
            )
            self.tables.append(table_item)
            self.blocks.append({"block_type": "table", **table_item})

    def _serialize_paragraph(self, paragraph: Any, paragraph_index: int) -> dict[str, Any]:
        text = _normalize_text(getattr(paragraph, "text", ""))
        style_name = ""
        try:
            style_name = str(paragraph.style.name or "")
        except Exception:
            style_name = ""

        return {
            "paragraph_index": paragraph_index,
            "text": text,
            "style_name": style_name,
            "level": _paragraph_level(text, style_name) if text else None,
        }

    def _infer_header_row(self, rows: list[dict[str, Any]]) -> tuple[int | None, list[str]]:
        if not rows:
            return None, []

        for row in rows[:3]:
            values = [str(cell.get("text") or "").strip() for cell in row["cells"] if str(cell.get("text") or "").strip()]
            if len(values) < 2:
                continue
            lowered = " ".join(values).lower()
            if any(keyword in lowered for keyword in TRF_HEADER_KEYWORDS):
                return int(row["row_index"]), values

        first_row = rows[0]
        values = [str(cell.get("text") or "").strip() for cell in first_row["cells"] if str(cell.get("text") or "").strip()]
        if len(values) >= 2 and sum(len(value) for value in values) <= 240:
            return int(first_row["row_index"]), values

        return None, []

    def _infer_table_type(self, rows: list[dict[str, Any]], column_headers: list[str]) -> str:
        if not rows:
            return "empty_table"

        header_text = " ".join(column_headers).lower()
        if header_text and any(keyword in header_text for keyword in TRF_HEADER_KEYWORDS):
            return "trf_checklist"

        max_columns = max((len(row["cells"]) for row in rows), default=0)
        if max_columns <= 2:
            return "key_value"

        return "grid"

    def _build_table_title(
        self,
        rows: list[dict[str, Any]],
        table_index: int,
        preceding_heading: dict[str, Any] | None,
        preceding_paragraph: str | None,
    ) -> str:
        if preceding_paragraph and _looks_like_caption(preceding_paragraph):
            return _truncate_title(preceding_paragraph)

        if preceding_heading and str(preceding_heading.get("title") or "").strip():
            return f"Table {table_index}: {_truncate_title(str(preceding_heading['title']))}"

        first_non_empty_row = next((row for row in rows if row["text"]), None)
        if first_non_empty_row is not None:
            return f"Table {table_index}: {_truncate_title(str(first_non_empty_row['text']))}"

        return f"Table {table_index}"

    def _serialize_table(
        self,
        table: Any,
        *,
        table_index: int,
        preceding_heading: dict[str, Any] | None,
        preceding_paragraph: str | None,
        preceding_paragraph_index: int | None,
    ) -> dict[str, Any]:
        rows: list[dict[str, Any]] = []
        column_count = 0

        for row_index, row in enumerate(getattr(table, "rows", []), start=1):
            seen_cells: set[int] = set()
            cells: list[dict[str, Any]] = []

            for cell in getattr(row, "cells", []):
                tc = getattr(cell, "_tc", None)
                cell_id = id(tc) if tc is not None else id(cell)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                cells.append(
                    {
                        "column_index": len(cells) + 1,
                        "text": _normalize_text(getattr(cell, "text", "")),
                    }
                )

            column_count = max(column_count, len(cells))
            rows.append(
                {
                    "row_index": row_index,
                    "cells": cells,
                    "text": _format_row_text(cells),
                }
            )

        header_row_index, column_headers = self._infer_header_row(rows)
        table_type = self._infer_table_type(rows, column_headers)
        for row in rows:
            row["text"] = _format_row_text(row["cells"], table_type=table_type)

        return {
            "table_index": table_index,
            "title": self._build_table_title(rows, table_index, preceding_heading, preceding_paragraph),
            "table_type": table_type,
            "row_count": len(rows),
            "column_count": column_count,
            "header_row_index": header_row_index,
            "column_headers": column_headers,
            "rows": rows,
            "heading_title": str(preceding_heading.get("title") or "") if preceding_heading else "",
            "heading_level": int(preceding_heading.get("level") or 0) if preceding_heading else 0,
            "preceding_paragraph": preceding_paragraph or "",
            "preceding_paragraph_index": preceding_paragraph_index,
        }

    @property
    def paragraph_count(self) -> int:
        return len(self.paragraphs)

    @property
    def table_count(self) -> int:
        return len(self.tables)

    def _build_table_nodes(self, max_nodes: int) -> list[dict[str, Any]]:
        nodes: list[dict[str, Any]] = []
        for table in self.tables[:max_nodes]:
            nodes.append(
                {
                    "title": table["title"],
                    "level": max(1, int(table.get("heading_level") or 0) or 1),
                    "anchor": f"table-{int(table['table_index'])}",
                    "locator": {
                        "table_index": int(table["table_index"]),
                        "row_start": 1,
                        "row_end": int(table["row_count"]),
                    },
                    "table_type": table["table_type"],
                }
            )
        return nodes

    def get_structure(self, max_nodes: int = 200) -> dict[str, Any]:
        headings = [item for item in self.paragraphs if item["text"] and item["level"] is not None]

        if headings:
            nodes: list[dict[str, Any]] = []
            limited_headings = headings[:max_nodes]
            for index, item in enumerate(limited_headings):
                next_start = (
                    limited_headings[index + 1]["paragraph_index"]
                    if index + 1 < len(limited_headings)
                    else self.paragraph_count + 1
                )
                title = item["text"]
                nodes.append(
                    {
                        "title": title,
                        "level": int(item["level"] or 1),
                        "anchor": _make_anchor(title),
                        "locator": {
                            "paragraph_start": item["paragraph_index"],
                            "paragraph_end": max(item["paragraph_index"], next_start - 1),
                        },
                    }
                )
            structure_type = "word_heading_map"
        elif self.tables:
            nodes = self._build_table_nodes(max_nodes)
            structure_type = "word_table_map"
        else:
            paragraphs = [item for item in self.paragraphs if item["text"]]
            nodes = [
                {
                    "title": _truncate_title(item["text"]),
                    "level": 1,
                    "anchor": f"paragraph-{item['paragraph_index']}",
                    "locator": {
                        "paragraph_start": item["paragraph_index"],
                        "paragraph_end": item["paragraph_index"],
                    },
                }
                for item in paragraphs[:max_nodes]
            ]
            structure_type = "word_paragraph_map"

        return {
            "docx_path": str(self.docx_path),
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
            "structure_type": structure_type,
            "items": nodes[:max_nodes],
        }

    def search(
        self,
        query: str,
        *,
        mode: str = "plain",
        case_sensitive: bool = False,
        max_results: int = 50,
        context_paragraphs: int = 2,
    ) -> dict[str, Any]:
        flags = 0 if case_sensitive else re.IGNORECASE
        pattern = re.compile(query if mode == "regex" else re.escape(query), flags)
        items: list[dict[str, Any]] = []
        non_empty_paragraphs = [item for item in self.paragraphs if item["text"]]

        for index, item in enumerate(non_empty_paragraphs):
            if len(items) >= max_results:
                break
            match = pattern.search(item["text"])
            if not match:
                continue
            start = max(0, index - context_paragraphs)
            end = min(len(non_empty_paragraphs), index + context_paragraphs + 1)
            context_before = "\n".join(entry["text"] for entry in non_empty_paragraphs[start:index] if entry["text"])
            context_after = "\n".join(entry["text"] for entry in non_empty_paragraphs[index + 1:end] if entry["text"])
            items.append(
                {
                    "source_type": "paragraph",
                    "paragraph_index": item["paragraph_index"],
                    "style_name": item["style_name"],
                    "match_text": match.group(0),
                    "text": item["text"],
                    "context_before": context_before,
                    "context_after": context_after,
                }
            )

        if len(items) < max_results:
            for table in self.tables:
                non_empty_rows = [row for row in table["rows"] if row["text"]]
                for row_position, row in enumerate(non_empty_rows):
                    if len(items) >= max_results:
                        break

                    for cell in row["cells"]:
                        cell_text = str(cell.get("text") or "")
                        if not cell_text:
                            continue
                        match = pattern.search(cell_text)
                        if not match:
                            continue

                        start = max(0, row_position - context_paragraphs)
                        end = min(len(non_empty_rows), row_position + context_paragraphs + 1)
                        context_before = "\n".join(entry["text"] for entry in non_empty_rows[start:row_position] if entry["text"])
                        context_after = "\n".join(entry["text"] for entry in non_empty_rows[row_position + 1:end] if entry["text"])
                        column_headers = list(table.get("column_headers") or [])
                        column_index = int(cell["column_index"])
                        column_header = ""
                        if column_headers and column_index <= len(column_headers):
                            column_header = str(column_headers[column_index - 1] or "")

                        items.append(
                            {
                                "source_type": "table_cell",
                                "table_index": int(table["table_index"]),
                                "row_index": int(row["row_index"]),
                                "column_index": column_index,
                                "table_title": str(table["title"]),
                                "table_type": str(table["table_type"]),
                                "section_title": str(table.get("heading_title") or ""),
                                "column_header": column_header,
                                "row_text": str(row["text"]),
                                "match_text": match.group(0),
                                "text": cell_text,
                                "context_before": context_before,
                                "context_after": context_after,
                            }
                        )
                        if len(items) >= max_results:
                            break

        return {
            "docx_path": str(self.docx_path),
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
            "items": items,
        }

    def read_paragraphs(
        self,
        paragraph_start: int,
        paragraph_end: int,
        *,
        include_context: int = 0,
    ) -> dict[str, Any]:
        if paragraph_start < 1 or paragraph_end < paragraph_start:
            raise ValueError("Invalid range: paragraph_end must be >= paragraph_start and paragraph_start must be >= 1")
        if paragraph_end > self.paragraph_count:
            raise ValueError(f"Paragraph range {paragraph_start}-{paragraph_end} exceeds paragraph count {self.paragraph_count}")

        requested = set(range(paragraph_start, paragraph_end + 1))
        selected = set(requested)
        if include_context > 0:
            selected.update(
                range(
                    max(1, paragraph_start - include_context),
                    min(self.paragraph_count, paragraph_end + include_context) + 1,
                )
            )

        items = []
        for item in self.paragraphs:
            if item["paragraph_index"] not in selected:
                continue
            items.append(
                {
                    "paragraph_index": item["paragraph_index"],
                    "style_name": item["style_name"],
                    "text": item["text"],
                    "requested": item["paragraph_index"] in requested,
                }
            )

        return {
            "docx_path": str(self.docx_path),
            "paragraph_count": self.paragraph_count,
            "paragraph_start": paragraph_start,
            "paragraph_end": paragraph_end,
            "include_context": include_context,
            "items": items,
        }

    def read_table_rows(
        self,
        table_index: int,
        *,
        row_start: int | None = None,
        row_end: int | None = None,
        column_start: int | None = None,
        column_end: int | None = None,
        include_context: int = 0,
    ) -> dict[str, Any]:
        if table_index < 1 or table_index > self.table_count:
            raise ValueError(f"Table index {table_index} is out of range (table count: {self.table_count})")

        table = self.tables[table_index - 1]
        total_rows = int(table["row_count"])
        total_columns = int(table["column_count"])

        normalized_row_start = 1 if row_start is None else int(row_start)
        normalized_row_end = total_rows if row_end is None else int(row_end)
        normalized_column_start = 1 if column_start is None else int(column_start)
        normalized_column_end = total_columns if column_end is None else int(column_end)

        if normalized_row_start < 1 or normalized_row_end < normalized_row_start:
            raise ValueError("Invalid row range: row_end must be >= row_start and row_start must be >= 1")
        if normalized_row_end > total_rows:
            raise ValueError(
                f"Row range {normalized_row_start}-{normalized_row_end} exceeds row count {total_rows}"
            )
        if normalized_column_start < 1 or normalized_column_end < normalized_column_start:
            raise ValueError(
                "Invalid column range: column_end must be >= column_start and column_start must be >= 1"
            )
        if normalized_column_end > total_columns:
            raise ValueError(
                f"Column range {normalized_column_start}-{normalized_column_end} exceeds column count {total_columns}"
            )

        requested = set(range(normalized_row_start, normalized_row_end + 1))
        selected = set(requested)
        if include_context > 0:
            selected.update(
                range(
                    max(1, normalized_row_start - include_context),
                    min(total_rows, normalized_row_end + include_context) + 1,
                )
            )

        header_row_index = int(table["header_row_index"]) if table.get("header_row_index") else None
        if header_row_index is not None:
            selected.add(header_row_index)

        items = []
        for row in table["rows"]:
            row_index = int(row["row_index"])
            if row_index not in selected:
                continue

            filtered_cells = [
                cell
                for cell in row["cells"]
                if normalized_column_start <= int(cell["column_index"]) <= normalized_column_end
            ]
            items.append(
                {
                    "row_index": row_index,
                    "requested": row_index in requested,
                    "is_header": row_index == header_row_index,
                    "cells": filtered_cells,
                    "text": _format_row_text(filtered_cells, table_type=str(table["table_type"])),
                }
            )

        return {
            "docx_path": str(self.docx_path),
            "table_index": table_index,
            "title": str(table["title"]),
            "table_type": str(table["table_type"]),
            "row_count": total_rows,
            "column_count": total_columns,
            "row_start": normalized_row_start,
            "row_end": normalized_row_end,
            "column_start": normalized_column_start,
            "column_end": normalized_column_end,
            "include_context": include_context,
            "header_row_index": header_row_index,
            "column_headers": list(table.get("column_headers") or []),
            "items": items,
        }


def get_word_structure(docx_path: str | Path, max_nodes: int = 200) -> dict[str, Any]:
    reader = WordReader(docx_path)
    return reader.get_structure(max_nodes=max_nodes)


def search_word_document(
    docx_path: str | Path,
    query: str,
    *,
    mode: str = "plain",
    case_sensitive: bool = False,
    max_results: int = 50,
    context_paragraphs: int = 2,
) -> dict[str, Any]:
    reader = WordReader(docx_path)
    return reader.search(
        query,
        mode=mode,
        case_sensitive=case_sensitive,
        max_results=max_results,
        context_paragraphs=context_paragraphs,
    )


def read_word_paragraphs(
    docx_path: str | Path,
    paragraph_start: int,
    paragraph_end: int,
    *,
    include_context: int = 0,
) -> dict[str, Any]:
    reader = WordReader(docx_path)
    return reader.read_paragraphs(
        paragraph_start,
        paragraph_end,
        include_context=include_context,
    )


def read_word_table_rows(
    docx_path: str | Path,
    table_index: int,
    *,
    row_start: int | None = None,
    row_end: int | None = None,
    column_start: int | None = None,
    column_end: int | None = None,
    include_context: int = 0,
) -> dict[str, Any]:
    reader = WordReader(docx_path)
    return reader.read_table_rows(
        table_index,
        row_start=row_start,
        row_end=row_end,
        column_start=column_start,
        column_end=column_end,
        include_context=include_context,
    )
